# -*- coding: utf-8 -*-
"""Génère le DOCX Chapitre 3 + Partie II (Analyse et conception) — FluxMin."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parents[1] / (
    "Chapitre3_PartieII_Analyse_Conception_FluxMin.docx"
)


def set_run_font(run, size=11, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12 if level == 2 else 11, bold=True)
    return p


def add_para(doc, text, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_code(doc, code: str, title: str | None = None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, size=10, bold=True)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(code.strip() + "\n")
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)
    # fond gris léger
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    shd.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shd)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, size=9, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9)
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # ========== PAGE DE GARDE SECTION ==========
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "Conception et réalisation d’une plateforme web sécurisée "
        "de gestion et d’automatisation intelligente des courriers "
        "et communications interministériels"
    )
    set_run_font(r, size=14, bold=True)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = st.add_run(
        "Mémoire de fin d’études — Master Informatique\n"
        "Chapitre 3 — Description du projet\n"
        "PARTIE II — Analyse et conception (Chapitres 4 et 5)\n"
        "Application : FluxMin"
    )
    set_run_font(r2, size=11)
    doc.add_paragraph()

    note = doc.add_paragraph()
    nr = note.add_run(
        "Note méthodologique : les diagrammes UML ne sont pas intégrés en image. "
        "Chaque section concernée fournit un bloc PlantUML à copier-coller "
        "(PlantUML Online, VS Code PlantUML, ou IntelliJ) pour générer la figure."
    )
    set_run_font(nr, size=10, italic=True)
    doc.add_page_break()

    # ========== CHAPITRE 3 ==========
    add_heading_custom(doc, "Chapitre 3. Description du projet", 1)

    add_heading_custom(doc, "3.1. Formulation", 2)
    add_para(
        doc,
        "Dans l’administration publique, la circulation des courriers et des "
        "communications officielles entre ministères demeure souvent fragmentée : "
        "mélange de courrier papier, de messagerie électronique non structurée et "
        "de procédures manuelles. Ces pratiques engendrent des pertes de temps, "
        "une traçabilité insuffisante, des risques d’erreur et une difficulté à "
        "piloter les délais de traitement.",
    )
    add_para(
        doc,
        "Le présent projet, intitulé FluxMin, vise la conception et la réalisation "
        "d’une plateforme web sécurisée permettant de digitaliser, fluidifier et "
        "automatiser les flux de courriers interministériels et intra-ministériels, "
        "tout en intégrant un canal de communications gouvernementales, une couche "
        "d’assistance par intelligence artificielle (OCR, résumé, suggestions) et "
        "des mécanismes d’hyperautomation (relances, escalades, analyse des flux).",
    )
    add_para(
        doc,
        "La problématique peut être formulée ainsi : comment concevoir et réaliser "
        "une plateforme logicielle souveraine, multi-ministères, capable d’assurer "
        "la gestion complète du cycle de vie des courriers, la communication "
        "officielle ciblée, la traçabilité d’audit et l’assistance intelligente, "
        "tout en respectant une organisation en directions et un contrôle d’accès "
        "fondé sur les rôles ?",
    )

    add_heading_custom(doc, "3.2. Objectif et besoins de l’utilisateur", 2)
    add_para(doc, "Objectif général", bold=True)
    add_para(
        doc,
        "Mettre à disposition des acteurs ministériels une plateforme unique pour "
        "créer, transmettre, recevoir, discuter, archiver et piloter les courriers, "
        "ainsi que pour diffuser et accuser réception des communications du "
        "Gouvernement, avec un gain de productivité et une traçabilité complète.",
    )
    add_para(doc, "Objectifs spécifiques", bold=True)
    for item in [
        "Digitaliser le circuit courrier interne (même ministère) et externe (entre ministères).",
        "Garantir l’authentification, l’autorisation RBAC et l’audit des actions sensibles.",
        "Permettre le multi-fichiers (pièces jointes) avec téléchargement et conservation.",
        "Notifier les utilisateurs en temps réel (WebSocket) des événements métier.",
        "Assister les agents via IA (OCR local, résumé/infos clés, suggestions validées humainement).",
        "Diffuser des actualités gouvernementales publiques ou ciblées, avec AR du directeur de ministère.",
        "Automatiser les relances/escalades et offrir un tableau d’analyse des flux (process mining).",
        "Documenter et durcir l’API (Swagger, rate limiting, health check).",
    ]:
        add_para(doc, f"• {item}")

    add_para(doc, "Besoins utilisateurs (synthèse)", bold=True)
    add_para(
        doc,
        "Les agents et responsables ont besoin d’un espace de travail simple pour "
        "traiter les courriers ; les agents courrier et directions destinataires "
        "doivent transmettre en interne ; les auditeurs consultent les journaux et "
        "anomalies ; le compte Gouvernement publie des communications ; le directeur "
        "de ministère accuse réception des posts ciblés ; le super administrateur "
        "configure ministères, directions et comptes sans intervenir dans le métier "
        "opérationnel des publications.",
    )

    add_heading_custom(
        doc,
        "3.3. Moyens nécessaires à la réalisation du projet (Humain, Matériel, Logiciel)",
        2,
    )
    add_para(doc, "Moyens humains", bold=True)
    add_table(
        doc,
        ["Rôle", "Responsabilités", "Implication"],
        [
            ["Stagiaire / développeur full-stack", "Conception, réalisation, tests, documentation", "Temps plein projet"],
            ["Encadrant pédagogique / maître de stage", "Orientation, validation des livrables", "Suivi périodique"],
            ["Product Owner (métier / tutelle)", "Priorisation du backlog, validation fonctionnelle", "Ateliers Scrum"],
            ["Utilisateurs pilotes (démo)", "Retours d’usage sur parcours courrier et actualités", "Sessions de test"],
        ],
    )
    add_para(doc, "Moyens matériels", bold=True)
    add_table(
        doc,
        ["Élément", "Caractéristiques / usage"],
        [
            ["Poste de développement", "PC Windows, IDE (Cursor/VS Code), navigateur moderne"],
            ["Serveur de base de données", "PostgreSQL 16 local (ou instance dédiée)"],
            ["Infrastructure optionnelle", "Conteneurs Docker pour MinIO, Temporal, Redis (si disponible)"],
            ["Réseau", "Accès local HTTP (ports 3000 frontend, 3001 API, 8000 IA)"],
        ],
    )
    add_para(doc, "Moyens logiciels", bold=True)
    add_table(
        doc,
        ["Couche", "Technologies"],
        [
            ["Frontend", "Next.js 15, TypeScript, Tailwind CSS, shadcn/ui, TanStack Query, Zustand"],
            ["Backend", "NestJS, TypeScript, Drizzle ORM, JWT, WebSocket (Socket.IO)"],
            ["Base de données", "PostgreSQL 16"],
            ["IA", "Python FastAPI, Tesseract/OCR local, LLM multi-fournisseurs (résumé)"],
            ["Stockage / workflow", "MinIO (S3), Temporal.io (relances), disque local en fallback"],
            ["Qualité / docs API", "Jest, Swagger/OpenAPI, Helmet, rate limiting"],
            ["Conception", "UML (PlantUML), méthode agile Scrum"],
        ],
    )

    add_heading_custom(doc, "3.4. Résultats attendus", 2)
    for item in [
        "Une application web opérationnelle FluxMin (frontend + backend + service IA).",
        "Un modèle organisationnel ministères / directions / utilisateurs avec RBAC.",
        "Le cycle de vie courrier : brouillon, envoi, réception, transmission, discussion, archivage.",
        "Le module Communications Gouvernement (publications, PJ, notifications, AR, fil de discussion).",
        "Des indicateurs de pilotage (dashboard, analytics, process mining).",
        "Une documentation de conception (présent document) et une API documentée (Swagger).",
        "Un jeu de comptes de démonstration et un protocole de tests fonctionnels.",
    ]:
        add_para(doc, f"• {item}")

    add_heading_custom(doc, "3.5. Chronogramme de travail", 2)
    add_para(
        doc,
        "Le projet a été conduit de façon itérative (modules M1 à M8 + MG). "
        "Le chronogramme ci-dessous synthétise les phases. "
        "Pour générer un diagramme de Gantt PlantUML, utiliser le code suivant.",
    )
    add_table(
        doc,
        ["Phase", "Contenu", "Période indicative"],
        [
            ["Initialisation", "Schéma DB, monorepo, auth, admin", "Semaines 1–2"],
            ["Cœur métier", "Courriers, flux, PJ, messagerie, notifications", "Semaines 3–5"],
            ["Pilotage & audit", "Archives, rétention, audit, anomalies", "Semaines 5–6"],
            ["Pilotage & IA", "Dashboard, analytics, OCR/LLM, suggestions", "Semaines 6–8"],
            ["Gouvernement (MG)", "Publications, AR directeur, discussions", "Semaine 8–9"],
            ["Hyperautomation (M7)", "MinIO, Temporal, process mining", "Semaine 9–10"],
            ["Durcissement (M8)", "Swagger, rate limit, health, tests", "Semaine 10–11"],
            ["Rédaction mémoire", "Analyse, conception, bilan", "En parallèle / fin"],
        ],
    )
    add_code(
        doc,
        r"""
@startgantt
title Chronogramme FluxMin — Master Informatique
project starts 2026-05-01
[Initialisation auth-admin] lasts 14 days
[Coeur courriers-messagerie] lasts 21 days
[Archivage et audit] lasts 14 days
[Pilotage et IA] lasts 21 days
[Communications Gouvernement] lasts 10 days
[Hyperautomation M7] lasts 10 days
[Durcissement M8] lasts 7 days
[Redaction memoire] lasts 30 days
[Initialisation auth-admin] -> [Coeur courriers-messagerie]
[Coeur courriers-messagerie] -> [Archivage et audit]
[Archivage et audit] -> [Pilotage et IA]
[Pilotage et IA] -> [Communications Gouvernement]
[Communications Gouvernement] -> [Hyperautomation M7]
[Hyperautomation M7] -> [Durcissement M8]
@endgantt
""",
        title="PlantUML — Chronogramme (à générer)",
    )

    doc.add_page_break()

    # ========== PARTIE II ==========
    add_heading_custom(doc, "PARTIE II. ANALYSE ET CONCEPTION", 1)

    add_heading_custom(doc, "Chapitre 4. Analyse préalable", 1)

    add_heading_custom(doc, "4.1. Analyse de l’existant", 2)
    add_heading_custom(
        doc, "4.1.1. Organisation actuelle (traitement actuel et personnel impliqué)", 3
    )
    add_para(
        doc,
        "Dans l’organisation administrative cible, chaque ministère est structuré en "
        "directions (Courrier, DSI, DRH, DAF, etc.). Les courriers transitent entre "
        "directions du même ministère (flux interne) ou d’un ministère vers une "
        "direction d’un autre ministère (flux externe). Le personnel impliqué comprend "
        "typiquement : agents de bureau, responsables de direction, agents du bureau "
        "du courrier, auditeurs, et, pour les communications d’État, un canal "
        "Gouvernement et un directeur par ministère.",
    )
    add_para(
        doc,
        "Avant FluxMin, le traitement courant repose souvent sur : enregistrement "
        "manuel ou tableurs, envoi par e-mail, classement papier, et absence de "
        "tableau de bord unifié des délais. Les accusés de réception des instructions "
        "gouvernementales sont difficiles à centraliser.",
    )
    add_code(
        doc,
        r"""
@startuml
title Organisation actuelle — acteurs du traitement courrier
left to right direction
actor "Agent" as A
actor "Responsable\ndirection" as R
actor "Agent courrier" as C
actor "Auditeur" as AU
actor "Directeur\nministere" as D
actor "Gouvernement" as G
rectangle "Ministere" {
  usecase "Reception papier / email" as U1
  usecase "Enregistrement manuel" as U2
  usecase "Transmission interne" as U3
  usecase "Classement / archive" as U4
}
A --> U1
C --> U2
R --> U3
AU --> U4
G ..> D : instruction officielle\n(hors systeme)
@enduml
""",
        title="PlantUML — Organisation / cas d’usage de l’existant (optionnel)",
    )

    add_heading_custom(doc, "4.1.2. Inventaire des moyens matériels et logiciels", 3)
    add_table(
        doc,
        ["Catégorie", "Existant typique"],
        [
            ["Matériel", "Postes bureautiques, imprimantes, scanners, serveurs de messagerie"],
            ["Logiciel bureautique", "Suite Office, clients mail (Outlook, webmail)"],
            ["Stockage", "Dossiers partagés réseau, disques locaux, parfois GED partielle"],
            ["Sécurité", "Comptes AD/LDAP éventuels, peu de RBAC métier fine sur le courrier"],
            ["Pilotage", "Reporting manuel ou Excel"],
        ],
    )

    add_heading_custom(doc, "4.2. Critique de l’existant (points forts et points faibles)", 2)
    add_para(doc, "Points forts", bold=True)
    for item in [
        "Culture administrative établie autour du courrier et de la hiérarchie.",
        "Présence de scanners et de compétences bureautiques.",
        "Besoin métier clairement exprimé (traçabilité, délais, formalisme).",
    ]:
        add_para(doc, f"• {item}")
    add_para(doc, "Points faibles", bold=True)
    for item in [
        "Absence de référentiel unique des courriers et de leur historique.",
        "Multiplicité des canaux (papier, mail) sans circuit normé.",
        "Difficulté à mesurer les délais et à détecter les anomalies de workflow.",
        "Pas d’assistance à l’extraction d’information depuis les pièces jointes.",
        "Communications gouvernementales peu instrumentées (AR, ciblage ministère).",
        "Risques de perte de pièces et de non-conformité d’archivage.",
    ]:
        add_para(doc, f"• {item}")

    add_heading_custom(doc, "4.3. Conception avant-projet", 2)
    add_heading_custom(doc, "4.3.1. Proposition des solutions", 3)
    add_para(
        doc,
        "Trois orientations ont été envisagées : (1) améliorer uniquement la messagerie "
        "existante ; (2) déployer une GED générique du marché ; (3) concevoir une "
        "plateforme métier dédiée (FluxMin) intégrant circuit courrier, RBAC "
        "ministériel, IA et communications gouvernementales. La solution (3) a été "
        "retenue pour coller au processus interministériel, garantir la souveraineté "
        "des données (stockage contrôlé, OCR local) et permettre l’évolution agile "
        "par modules.",
    )
    add_table(
        doc,
        ["Option", "Avantages", "Inconvénients", "Décision"],
        [
            ["Messagerie seule", "Coût faible", "Pas de circuit métier ni d’audit fin", "Rejetée"],
            ["GED générique", "Fonctions documentaires riches", "Adaptation métier lourde, licence", "Écartée"],
            ["Plateforme FluxMin", "Sur-mesure, IA, MG, traçabilité", "Effort de développement", "Retenue"],
        ],
    )

    add_heading_custom(
        doc,
        "4.3.2. Méthodes de conception et outils utilisés (Choix, Justification et présentation)",
        3,
    )
    add_para(
        doc,
        "Méthode agile Scrum : livraisons incrémentales alignées sur les modules "
        "(M1–M8, MG), priorisation par valeur métier, démonstrations régulières. "
        "La modélisation s’appuie sur UML (diagrammes de classes, paquetages, "
        "déploiement, modèle de domaine) générés via PlantUML. "
        "Les choix techniques (NestJS, Next.js, PostgreSQL, FastAPI) privilégient "
        "la productivité TypeScript/Python, un écosystème mature et une séparation "
        "claire front / API / IA.",
    )

    doc.add_page_break()

    # ========== CHAPITRE 5 ==========
    add_heading_custom(doc, "Chapitre 5. Analyse conceptuelle", 1)

    add_heading_custom(doc, "5.1. Désignation des rôles de l’équipe Scrum", 2)
    add_para(
        doc,
        "Dans le contexte du stage / mémoire, les rôles Scrum sont adaptés à une "
        "équipe réduite.",
    )
    add_table(
        doc,
        ["Rôle Scrum", "Affectation dans le contexte du stage", "Responsabilités"],
        [
            [
                "Product Owner",
                "Encadrant métier / tuteur institutionnel (ou stagiaire en double casquette avec validation tutelle)",
                "Vision produit, priorisation du backlog, acceptance des user stories",
            ],
            [
                "Scrum Master",
                "Stagiaire (facilitation) assisté de l’encadrant pédagogique",
                "Animation des rituels, levée des obstacles, respect du cadre agile",
            ],
            [
                "Development Team",
                "Stagiaire développeur full-stack (éventuellement binôme ponctuel)",
                "Conception, code, tests, intégration front/back/IA, documentation",
            ],
        ],
    )

    add_heading_custom(doc, "5.2. Élaboration du Product Backlog", 2)
    add_heading_custom(doc, "5.2.1. Vision du projet", 3)
    add_para(
        doc,
        "FluxMin devient la plateforme de référence pour numériser les échanges "
        "documentaires entre ministères : un agent traite un courrier de bout en "
        "bout ; le Gouvernement diffuse une instruction ciblée ; le directeur accuse "
        "réception ; l’auditeur vérifie la conformité ; l’IA accélère la lecture "
        "des pièces jointes sans se substituer à la décision humaine.",
    )

    add_heading_custom(doc, "5.2.2. Listes des acteurs", 3)
    add_table(
        doc,
        ["Acteur", "Rôle", "Description", "Fréquence d’utilisation"],
        [
            ["Super Admin", "Administration plateforme", "Gère ministères, directions, utilisateurs ; pas de messagerie métier", "Faible (configuration)"],
            ["Agent / Responsable", "Traitement courrier", "Crée, envoie, reçoit, transmet, discute, archive", "Quotidienne"],
            ["Agent Courrier", "Point d’entrée courrier", "Réception / orientation, mêmes transmissions internes autorisées", "Quotidienne"],
            ["Responsable de direction", "Pilotage direction", "Suit les flux, analytics, validations selon permissions", "Quotidienne / hebdo"],
            ["Directeur de ministère", "Autorité ministère", "AR et réponses aux publications gouvernementales ciblées ; sans direction", "Hebdomadaire / événementielle"],
            ["Gouvernement", "Canal officiel État", "Publie actualités publiques ou ciblées, PJ, suit les échanges", "Événementielle"],
            ["Auditeur", "Contrôle", "Consulte audit, rapports, anomalies ; lecture transversale", "Hebdomadaire"],
            ["Service IA", "Système", "OCR, analyse, suggestions (acteur technique)", "À la demande"],
        ],
    )

    add_heading_custom(doc, "5.2.3. Thèmes (fonctionnalités et modules)", 3)
    add_table(
        doc,
        ["ID thème", "Thème / Module", "Fonctionnalités principales"],
        [
            ["T1", "Fondations & Admin", "Auth JWT, RBAC, ministères, directions, utilisateurs"],
            ["T2", "Courriers & flux", "CRUD, envoi, réception, transmission interne, historique"],
            ["T3", "Fichiers & messagerie", "PJ multi-fichiers, discussion temps réel, notifications"],
            ["T4", "Archivage & Audit", "Rétention, scopes, logs, anomalies, rapports"],
            ["T5", "Pilotage & IA", "Dashboard, analytics, OCR/LLM, suggestions"],
            ["T6", "Communications Gouvernement", "Publications, PJ, notifs, AR, discussion officielle"],
            ["T7", "Hyperautomation", "MinIO, Temporal relances/escalades, process mining"],
            ["T8", "Durcissement", "Swagger, rate limit, helmet, health, tests"],
        ],
    )

    add_heading_custom(doc, "5.2.4. Vision de la première release", 3)
    add_para(
        doc,
        "Ordre d’importance des thèmes : T1 > T2 > T3 > T4 > T5 > T6 > T7 > T8. "
        "La première release (Release 1 — MVP métier) regroupe T1, T2 et T3 : "
        "authentifier les utilisateurs, gérer l’organisation, traiter un courrier "
        "avec pièces jointes, discussion et notifications. Les releases suivantes "
        "ajoutent l’archivage/audit (R2), le pilotage/IA (R3), le canal gouvernement "
        "(R4), puis l’hyperautomation et le durcissement (R5).",
    )

    add_heading_custom(doc, "5.2.5. Users stories", 3)
    stories = [
        ("Super Admin", "En tant que super administrateur, je veux créer un ministère et ses directions afin d’organiser les utilisateurs selon la structure administrative."),
        ("Super Admin", "En tant que super administrateur, je veux créer un utilisateur et lui affecter un rôle afin de contrôler ses droits d’accès."),
        ("Agent", "En tant qu’agent, je veux créer un courrier brouillon avec objet et destinataire afin de préparer un envoi officiel."),
        ("Agent", "En tant qu’agent, je veux joindre plusieurs fichiers à un courrier afin de transmettre les pièces justificatives."),
        ("Agent", "En tant qu’agent, je veux envoyer un courrier afin que la direction destinataire en soit notifiée."),
        ("Responsable", "En tant que responsable de direction, je veux transmettre un courrier à une autre direction du même ministère afin de poursuivre le traitement interne."),
        ("Agent", "En tant qu’agent, je veux discuter sur un courrier en temps réel afin de clarifier le dossier avec l’interlocuteur."),
        ("Agent", "En tant qu’agent, je veux archiver un courrier avec une durée de conservation afin de respecter la politique de rétention."),
        ("Auditeur", "En tant qu’auditeur, je veux consulter les journaux d’audit et les anomalies afin de contrôler la conformité des traitements."),
        ("Responsable", "En tant que responsable, je veux consulter le dashboard et les analytics afin de piloter l’activité de ma direction."),
        ("Agent", "En tant qu’agent, je veux lancer une analyse IA sur les pièces jointes afin d’obtenir un résumé et des suggestions que je valide."),
        ("Gouvernement", "En tant que compte Gouvernement, je veux publier une actualité ciblée à un ministère afin de diffuser une instruction officielle."),
        ("Directeur", "En tant que directeur de ministère, je veux accuser réception d’une publication ciblée afin de confirmer officiellement la prise en charge."),
        ("Directeur", "En tant que directeur de ministère, je veux échanger dans le fil de discussion d’une publication afin de répondre au Gouvernement."),
        ("Système", "En tant que système, je veux déclencher une relance puis une escalade si un courrier reste sans action afin de réduire les délais de traitement."),
    ]
    for actor, text in stories:
        add_para(doc, f"[{actor}] {text}")

    add_heading_custom(doc, "5.3. Product Backlog", 2)
    add_para(
        doc,
        "Le product backlog ci-dessous synthétise les items prioritaires. "
        "L’estimation est en jours-homme pour une équipe réduite.",
    )
    add_table(
        doc,
        ["ID_item", "Titre", "Importance", "Estimation (j)", "Démonstration", "Commentaire"],
        [
            ["PB01", "Authentification JWT + refresh", "100", "3", "Login, refresh, profil", "Sécurité de base"],
            ["PB02", "Admin ministères / directions / users", "95", "4", "CRUD + RBAC", ""],
            ["PB03", "Gestion des courriers (CRUD)", "90", "5", "Création / détail", ""],
            ["PB04", "Envoi / réception / transmission", "90", "5", "Circuit interne/externe", "Transmission interne uniquement"],
            ["PB05", "Pièces jointes multi-fichiers", "85", "4", "Upload / download", "Fallback disque si MinIO down"],
            ["PB06", "Messagerie temps réel", "80", "4", "Chat + présence", "WebSocket JWT"],
            ["PB07", "Notifications in-app", "80", "3", "Toast + liste", ""],
            ["PB08", "Archivage + rétention", "75", "4", "Archiver / désarchiver", ""],
            ["PB09", "Audit + anomalies + rapports", "70", "5", "Search / export", ""],
            ["PB10", "Dashboard + analytics", "70", "4", "Stats réelles", ""],
            ["PB11", "IA OCR + résumé + suggestions", "65", "6", "Analyser PJ", "Validation humaine"],
            ["PB12", "Communications Gouvernement", "85", "5", "Publier + AR", "Directeur sans direction"],
            ["PB13", "MinIO + Temporal + process mining", "60", "5", "Stockage / relance / flux", "Optionnel sans Docker"],
            ["PB14", "Durcissement API (Swagger, throttle)", "55", "2", "/api/docs + /health", "M8"],
        ],
    )

    add_heading_custom(doc, "5.4. Sprint Backlog", 2)
    add_heading_custom(doc, "5.4.1. Tableau Sprint découpé par release", 3)
    add_table(
        doc,
        ["Release", "ID item", "Titre de sprint", "Estimation (j)"],
        [
            ["R1 MVP", "PB01–PB07", "Sprint 1 — Fondations & Auth", "7"],
            ["R1 MVP", "PB03–PB05", "Sprint 2 — Courriers & fichiers", "9"],
            ["R1 MVP", "PB06–PB07", "Sprint 3 — Messagerie & notifs", "7"],
            ["R2 Gouvernance", "PB08–PB09", "Sprint 4 — Archivage & audit", "9"],
            ["R3 Pilotage IA", "PB10–PB11", "Sprint 5 — Analytics & IA", "10"],
            ["R4 Gouvernement", "PB12", "Sprint 6 — Module MG", "5"],
            ["R5 Hyperauto & Qualité", "PB13–PB14", "Sprint 7 — M7 & M8", "7"],
        ],
    )

    add_heading_custom(doc, "5.4.2. Tableau planification des tâches par Sprint", 3)
    add_table(
        doc,
        ["Sprint", "ID tâche", "Tâche", "Estimation (j)"],
        [
            ["S1", "T01", "Schéma DB + seed + auth JWT", "2"],
            ["S1", "T02", "Guards RBAC + pages admin", "2"],
            ["S1", "T03", "Login UI + AuthGuard frontend", "2"],
            ["S1", "T04", "Tests manuels auth", "1"],
            ["S2", "T05", "API courriers CRUD + flux_etapes", "3"],
            ["S2", "T06", "Envoi / transmission / réception", "3"],
            ["S2", "T07", "Upload PJ + UI détail", "2"],
            ["S2", "T08", "Pages inbox/sent/drafts", "1"],
            ["S3", "T09", "Gateway messages + présence", "2"],
            ["S3", "T10", "UI discussion + PJ message", "2"],
            ["S3", "T11", "Notifications WS + page liste", "2"],
            ["S3", "T12", "Recette R1", "1"],
            ["S4", "T13", "Archivage rétention scopes", "2"],
            ["S4", "T14", "Audit interceptor + UI search", "3"],
            ["S4", "T15", "Anomalies + rapports export", "3"],
            ["S4", "T16", "Recette R2", "1"],
            ["S5", "T17", "Stats dashboard/analytics", "2"],
            ["S5", "T18", "Service IA FastAPI OCR", "3"],
            ["S5", "T19", "Suggestions + validation UI", "3"],
            ["S5", "T20", "Recette R3", "1"],
            ["S6", "T21", "Schéma publications + API", "2"],
            ["S6", "T22", "UI publier / actualités / AR", "2"],
            ["S6", "T23", "Discussion posts + notifs", "1"],
            ["S7", "T24", "Storage MinIO + dual-read", "2"],
            ["S7", "T25", "Workflows Temporal + mining UI", "3"],
            ["S7", "T26", "Swagger, throttle, health, tests", "2"],
        ],
    )

    add_heading_custom(doc, "5.5. Dictionnaire des données", 2)
    add_para(
        doc,
        "Les données sont présentées par ordre alphabétique du libellé. "
        "Types : A (alphabétique), AN (alphanumérique), N (numérique), D (date/heure).",
    )
    dict_rows = [
        ["Action (flux)", "Type d’étape du circuit courrier (envoi, réception…)", "AN", "50", ""],
        ["Chemin fichier / MinIO", "Emplacement de stockage d’une pièce jointe", "AN", "512", "Disque ou clé objet"],
        ["Code ministère", "Sigle court du ministère (ex. MFA)", "AN", "50", "Unique"],
        ["Commentaire AR", "Texte optionnel d’accusé de réception publication", "A", "texte", ""],
        ["Corps", "Corps textuel du courrier ou de la publication", "A", "texte", ""],
        ["Date action", "Horodatage d’une étape de flux", "D", "—", "AAAA-MM-JJ hh:mm"],
        ["Date archivage", "Date d’entrée en archive", "D", "—", ""],
        ["Date envoi", "Date d’envoi du courrier", "D", "—", ""],
        ["Date publication", "Date de publication d’une actualité", "D", "—", ""],
        ["Date réception", "Date d’accusé / réception courrier", "D", "—", ""],
        ["Direction (nom)", "Libellé de la direction", "A", "255", ""],
        ["Durée conservation", "Durée d’archivage en années", "N", "entier", "≥ 1"],
        ["Email", "Identifiant de connexion de l’utilisateur", "AN", "255", "Unique"],
        ["Emplacement archive", "Lieu physique/logique d’archivage", "AN", "255", "Optionnel"],
        ["Estimation / importance", "Métadonnées de backlog (hors runtime)", "N", "—", "Gestion de projet"],
        ["Hash SHA-256", "Empreinte d’intégrité du fichier", "AN", "64", "Optionnel"],
        ["Identifiant (id)", "Clé primaire technique", "N", "serial", ""],
        ["Message (contenu)", "Texte d’un message de discussion", "A", "texte", ""],
        ["Métadonnées IA", "JSON d’extraction OCR / analyse", "AN", "jsonb", ""],
        ["Ministère (nom)", "Nom officiel du ministère", "A", "255", "Unique"],
        ["Mot de passe", "Secret d’authentification (hashé)", "AN", "255", "Jamais en clair"],
        ["Nom / Prénom", "Identité de l’utilisateur", "A", "100", ""],
        ["Nom fichier", "Nom original de la pièce jointe", "AN", "255", ""],
        ["Notifications (titre, type)", "Alerte in-app", "AN", "var.", "lu: booléen"],
        ["Objet", "Objet du courrier", "A", "texte", "Obligatoire"],
        ["Permissions", "Droits fins éventuellement surchargés", "AN", "jsonb", ""],
        ["Portée publication", "public ou ministere", "A", "20", ""],
        ["Priorité publication", "normale, haute, urgente", "A", "20", ""],
        ["Référence courrier", "Référence unique métier", "AN", "100", "ex. MFA-2026-000123"],
        ["Rôle", "Profil RBAC (agent_courrier, directeur_ministere…)", "A", "50", ""],
        ["Statut courrier", "brouillon, envoye, recu, en_traitement, archive", "A", "50", ""],
        ["Statut publication", "brouillon, publie, archive", "A", "50", ""],
        ["Taille octets", "Taille du fichier", "N", "bigint", ""],
        ["Titre publication", "Titre de l’actualité gouvernementale", "A", "255", ""],
        ["Type courrier", "interne ou externe", "A", "50", ""],
        ["Type MIME", "Type de fichier", "AN", "100", ""],
        ["Type publication", "communique, information, ordre, alerte", "A", "50", ""],
    ]
    add_table(
        doc,
        ["Information", "Description", "Type", "Taille", "Observation"],
        dict_rows,
    )

    add_heading_custom(doc, "5.6. Règles de gestion", 2)
    rules = [
        "RG01 — Tout utilisateur métier (hors super_admin / gouvernement selon cas) est rattaché soit à une direction, soit, pour le directeur_ministere, uniquement à un ministère (sans direction).",
        "RG02 — Un courrier interne circule entre directions du même ministère ; un courrier externe cible une direction d’un autre ministère à l’envoi.",
        "RG03 — Seule la direction actuellement destinataire peut transmettre le courrier, et uniquement vers une autre direction du même ministère.",
        "RG04 — Les pièces jointes respectent une liste de formats autorisés et une taille maximale (50 Mo courrier / publication).",
        "RG05 — L’archivage impose une durée de conservation minimale et met le courrier en statut archive ; la discussion devient lecture seule.",
        "RG06 — Toute mutation sensible est journalisée (audit).",
        "RG07 — Les suggestions IA sont proposées ; l’humain accepte, modifie ou ignore.",
        "RG08 — Une publication publique est visible de tous les utilisateurs métier ; une publication ciblée n’est visible que du ministère destinataire (et du Gouvernement).",
        "RG09 — Seul le directeur_ministere du ministère ciblé peut accuser réception et répondre sur une publication ciblée ; le super_admin ne répond pas aux publications.",
        "RG10 — Les relances/escalades Temporal ne s’appliquent qu’aux courriers encore en statut d’attente et s’annulent à l’AR ou à l’archivage.",
        "RG11 — Les accès API sont protégés par JWT ; le rate limiting limite les abus (notamment sur /auth/login).",
        "RG12 — Un seul compte gouvernement publie ; un directeur par ministère assure le canal de réponse officiel.",
    ]
    for r in rules:
        add_para(doc, r)

    add_heading_custom(doc, "5.7. Modèle du domaine", 2)
    add_para(
        doc,
        "Le modèle du domaine représente les concepts métier centraux et leurs "
        "relations (sans détail technique ORM). Générer le diagramme avec le code PlantUML suivant.",
    )
    add_code(
        doc,
        r"""
@startuml
title Modele du domaine — FluxMin
skinparam classAttributeIconSize 0

class Ministere
class Direction
class Utilisateur
class Courrier
class PieceJointe
class FluxEtape
class Archive
class Message
class Notification
class PublicationGouvernement
class AccuseReceptionPublication

Ministere "1" o-- "*" Direction : regroupe
Ministere "1" o-- "*" Utilisateur : rattache\n(directeur)
Direction "1" o-- "*" Utilisateur : emploie
Utilisateur "1" -- "*" Courrier : emet
Direction "1" -- "*" Courrier : emet / recoit
Courrier "1" o-- "*" PieceJointe
Courrier "1" o-- "*" FluxEtape
Courrier "0..1" -- "0..1" Archive
Courrier "1" o-- "*" Message
Utilisateur "1" -- "*" Notification : recoit
Utilisateur "1" -- "*" PublicationGouvernement : publie
Ministere "0..1" -- "*" PublicationGouvernement : cible
PublicationGouvernement "1" o-- "*" AccuseReceptionPublication
Utilisateur "1" -- "*" AccuseReceptionPublication : signe
@enduml
""",
        title="PlantUML — 5.7 Modèle du domaine",
    )

    add_heading_custom(doc, "5.8. Diagrammes de classes de conception par sprint", 2)
    add_para(
        doc,
        "Pour chaque sprint majeur, un diagramme de classes de conception (vue simplifiée "
        "des entités touchées) est proposé. Copier chaque bloc dans PlantUML.",
    )

    add_para(doc, "Sprint 1 — Fondations & Auth", bold=True)
    add_code(
        doc,
        r"""
@startuml
title Classes de conception — Sprint 1 (Auth / Admin)
class Ministere {
  +id: int
  +nom: string
  +code: string
}
class Direction {
  +id: int
  +nom: string
  +type: string
  +ministereId: int
}
class Utilisateur {
  +id: int
  +email: string
  +nom: string
  +prenom: string
  +role: string
  +permissions: json
  +motDePasse: string
  +directionId: int?
  +ministereId: int?
}
class AuthService
class AdminService
Ministere "1" --> "*" Direction
Direction "1" --> "*" Utilisateur
Ministere "1" --> "*" Utilisateur
AuthService ..> Utilisateur
AdminService ..> Ministere
AdminService ..> Direction
AdminService ..> Utilisateur
@enduml
""",
        title="PlantUML — Sprint 1",
    )

    add_para(doc, "Sprint 2 — Courriers & fichiers", bold=True)
    add_code(
        doc,
        r"""
@startuml
title Classes de conception — Sprint 2 (Courriers)
class Courrier {
  +id: int
  +reference: string
  +objet: string
  +corps: string
  +typeCourrier: string
  +statut: string
  +emetteurId: int
  +destinataireDirectionId: int
  +ministereDestinataireId: int?
}
class PieceJointe {
  +id: int
  +nomFichier: string
  +cheminMinio: string
  +typeMime: string
  +tailleBytes: long
}
class FluxEtape {
  +id: int
  +action: string
  +commentaire: string
  +dateAction: datetime
}
class CourrierService
class StorageService
Courrier "1" *-- "*" PieceJointe
Courrier "1" *-- "*" FluxEtape
CourrierService ..> Courrier
CourrierService ..> StorageService
@enduml
""",
        title="PlantUML — Sprint 2",
    )

    add_para(doc, "Sprint 3 — Messagerie & notifications", bold=True)
    add_code(
        doc,
        r"""
@startuml
title Classes de conception — Sprint 3
class Message {
  +id: int
  +contenu: string
  +courrierId: int
  +utilisateurId: int
}
class MessagePieceJointe {
  +id: int
  +cheminFichier: string
}
class Notification {
  +id: int
  +type: string
  +titre: string
  +lu: boolean
  +courrierId: int?
  +publicationId: int?
}
class MessagingService
class NotificationService
Message "1" *-- "*" MessagePieceJointe
MessagingService ..> Message
NotificationService ..> Notification
@enduml
""",
        title="PlantUML — Sprint 3",
    )

    add_para(doc, "Sprint 4 — Archivage & audit", bold=True)
    add_code(
        doc,
        r"""
@startuml
title Classes de conception — Sprint 4
class Archive {
  +id: int
  +courrierId: int
  +dateArchivage: datetime
  +dureeConservation: int
  +emplacement: string
}
class AuditLog {
  +id: int
  +utilisateurId: int
  +action: string
  +entiteType: string
  +entiteId: int
  +details: json
}
class ArchiveService
class AuditService
ArchiveService ..> Archive
AuditService ..> AuditLog
@enduml
""",
        title="PlantUML — Sprint 4",
    )

    add_para(doc, "Sprint 5 — Pilotage & IA", bold=True)
    add_code(
        doc,
        r"""
@startuml
title Classes de conception — Sprint 5
class StatsService
class AiService
class Courrier {
  +metadata: json
  +iaSuggestions: json
}
class PieceJointe {
  +metadataIa: json
}
AiService ..> Courrier
AiService ..> PieceJointe
StatsService ..> Courrier
@enduml
""",
        title="PlantUML — Sprint 5",
    )

    add_para(doc, "Sprint 6 — Communications Gouvernement", bold=True)
    add_code(
        doc,
        r"""
@startuml
title Classes de conception — Sprint 6 (MG)
class PublicationGouvernement {
  +id: int
  +titre: string
  +corps: string
  +portee: string
  +statut: string
  +ministereId: int?
  +auteurId: int
}
class PublicationPieceJointe
class PublicationMessage
class PublicationAccuse
class GouvernementService
PublicationGouvernement "1" *-- "*" PublicationPieceJointe
PublicationGouvernement "1" *-- "*" PublicationMessage
PublicationGouvernement "1" *-- "*" PublicationAccuse
GouvernementService ..> PublicationGouvernement
@enduml
""",
        title="PlantUML — Sprint 6",
    )

    add_para(doc, "Sprint 7 — Hyperautomation & durcissement", bold=True)
    add_code(
        doc,
        r"""
@startuml
title Classes de conception — Sprint 7
class StorageService
class TemporalService
class HealthService
StorageService : persistMulterFile()
StorageService : openReadStream()
TemporalService : startCourrierSuivi()
TemporalService : cancelCourrierSuivi()
HealthService : check()
@enduml
""",
        title="PlantUML — Sprint 7",
    )

    add_heading_custom(doc, "5.9. Diagramme de classes de conception globale", 2)
    add_para(
        doc,
        "Vue d’ensemble des principales classes persistantes et services applicatifs.",
    )
    add_code(
        doc,
        r"""
@startuml
title Diagramme de classes de conception globale — FluxMin
skinparam classAttributeIconSize 0

package "Organisation" {
  class Ministere
  class Direction
  class Utilisateur
}
package "Courrier" {
  class Courrier
  class PieceJointe
  class FluxEtape
  class Archive
  class Message
  class MessagePieceJointe
}
package "Gouvernement" {
  class PublicationGouvernement
  class PublicationPieceJointe
  class PublicationMessage
  class PublicationAccuse
}
package "Transverse" {
  class Notification
  class AuditLog
}
package "Services" {
  class AuthService
  class CourrierService
  class MessagingService
  class GouvernementService
  class ArchiveService
  class AiService
  class StorageService
  class TemporalService
  class StatsService
}

Ministere "1" *-- "*" Direction
Direction "1" o-- "*" Utilisateur
Ministere "1" o-- "0..*" Utilisateur
Utilisateur "1" --> "*" Courrier : emet
Courrier "*" --> "1" Direction : destinataire
Courrier "1" *-- "*" PieceJointe
Courrier "1" *-- "*" FluxEtape
Courrier "1" o-- "0..1" Archive
Courrier "1" *-- "*" Message
Message "1" *-- "*" MessagePieceJointe
Utilisateur "1" --> "*" PublicationGouvernement : auteur
PublicationGouvernement "*" --> "0..1" Ministere : cible
PublicationGouvernement "1" *-- "*" PublicationPieceJointe
PublicationGouvernement "1" *-- "*" PublicationMessage
PublicationGouvernement "1" *-- "*" PublicationAccuse
Utilisateur "1" --> "*" Notification
CourrierService ..> Courrier
CourrierService ..> StorageService
CourrierService ..> TemporalService
GouvernementService ..> PublicationGouvernement
MessagingService ..> Message
ArchiveService ..> Archive
AiService ..> PieceJointe
StatsService ..> Courrier
AuthService ..> Utilisateur
@enduml
""",
        title="PlantUML — 5.9 Conception globale",
    )

    add_heading_custom(doc, "5.10. Diagramme de paquetage", 2)
    add_code(
        doc,
        r"""
@startuml
title Diagramme de paquetages — FluxMin
skinparam packageStyle rectangle

package "frontend (Next.js)" {
  [App Router / Pages]
  [Composants UI]
  [Stores Zustand]
  [Client API]
}

package "backend (NestJS)" {
  [AuthModule]
  [AdminModule]
  [CourrierModule]
  [MessagingModule]
  [NotificationModule]
  [ArchiveModule]
  [AuditModule]
  [StatsModule]
  [AiModule]
  [GouvernementModule]
  [StorageModule]
  [TemporalModule]
}

package "ia-service (FastAPI)" {
  [OCR]
  [LLM Resume]
  [Analyse documents]
}

database "PostgreSQL" as PG
cloud "MinIO (optionnel)" as S3
cloud "Temporal (optionnel)" as TMP
cloud "LLM Cloud (optionnel)" as LLM

[Client API] --> [AuthModule]
[Client API] --> [CourrierModule]
[Client API] --> [GouvernementModule]
[AiModule] --> [OCR]
[AiModule] --> [LLM Resume]
[LLM Resume] ..> LLM : si cles configurees
[CourrierModule] --> PG
[StorageModule] --> S3
[StorageModule] --> PG
[TemporalModule] --> TMP
[TemporalModule] --> PG
@enduml
""",
        title="PlantUML — 5.10 Paquetages",
    )

    add_heading_custom(doc, "5.11. Diagramme de déploiement", 2)
    add_para(
        doc,
        "Le déploiement de développement local (sans Docker obligatoire) et la cible "
        "conteneurisée sont représentés ci-dessous.",
    )
    add_code(
        doc,
        r"""
@startuml
title Diagramme de deploiement — FluxMin (local + option Docker)
node "Poste developpeur / Serveur appli" {
  artifact "Frontend Next.js\n:3000" as FE
  artifact "Backend NestJS\n:3001/api" as BE
  artifact "IA FastAPI\n:8000" as IA
  folder "uploads/ (fallback disque)" as DISK
}
database "PostgreSQL\n:5432" as PG
node "Infrastructure optionnelle (Docker)" {
  artifact "MinIO\n:9000/:9001" as MINIO
  artifact "Temporal\n:7233" as TEMP
  artifact "Redis\n:6379" as REDIS
}
actor "Utilisateur navigateur" as U
U --> FE : HTTPS/HTTP
FE --> BE : REST + WS
BE --> PG
BE --> IA : HTTP
BE --> DISK
BE ..> MINIO : S3 si disponible
BE ..> TEMP : workflows si disponible
BE ..> REDIS : cache futur
@enduml
""",
        title="PlantUML — 5.11 Déploiement",
    )

    add_para(
        doc,
        "Fin du Chapitre 5. Les diagrammes générés depuis les blocs PlantUML "
        "ci-dessus sont à insérer dans la version finale du mémoire aux emplacements "
        "correspondants (légendes Figure 5.x).",
        italic=True,
    )

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
